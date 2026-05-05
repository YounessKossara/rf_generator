"""
DOCX Export — Professional Word report generator for OmniTest.

Generates a clean, presentable .docx report containing:
- Report title + execution metadata
- Summary table (passed/failed/total)
- Detailed test case results with screenshots
- Saved to /reports/ folder
"""

import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ── Color palette ──
COLOR_PASS = RGBColor(16, 185, 129)    # Green
COLOR_FAIL = RGBColor(239, 68, 68)     # Red
COLOR_MUTED = RGBColor(100, 116, 139)  # Slate
COLOR_TITLE = RGBColor(14, 165, 233)   # Marine blue


def create_test_report_docx(data: dict) -> bytes:
    """
    Generate a professional DOCX report from test results.
    
    Args:
        data: dict with keys: results, total_tests, passed, failed,
              target_url, browser, page_info, timestamp, etc.
    
    Returns:
        bytes of the generated .docx file
    """
    doc = Document()
    
    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════
    title = doc.add_heading("OmniTest — Rapport de Test", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = COLOR_TITLE

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Rapport généré automatiquement par l'agent QA autonome")
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # METADATA TABLE
    # ══════════════════════════════════════════
    doc.add_heading("Informations Générales", level=1)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    target_url = data.get("target_url", data.get("url", "N/A"))
    browser = data.get("browser", "chromium")
    total = data.get("total_tests", len(data.get("results", [])))
    passed = data.get("passed", 0)
    failed = data.get("failed", 0)

    meta_rows = [
        ("Date d'exécution", now),
        ("URL testée", target_url),
        ("Navigateur", _browser_label(browser)),
        ("Tests exécutés", str(total)),
        ("Réussis", str(passed)),
        ("Échoués", str(failed)),
        ("Taux de réussite", f"{(passed / total * 100):.0f}%" if total > 0 else "N/A"),
    ]

    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (key, val) in enumerate(meta_rows):
        row = table.rows[i]
        # Key cell (bold)
        cell_key = row.cells[0]
        cell_key.text = ""
        run = cell_key.paragraphs[0].add_run(key)
        run.bold = True
        run.font.size = Pt(10)
        # Value cell
        cell_val = row.cells[1]
        cell_val.text = val

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # RESULTS SUMMARY
    # ══════════════════════════════════════════
    doc.add_heading("Résumé des Résultats", level=1)

    pass_rate = (passed / total * 100) if total > 0 else 0
    summary_para = doc.add_paragraph()
    
    run_total = summary_para.add_run(f"Total: {total}  |  ")
    run_total.font.size = Pt(12)
    run_total.bold = True
    
    run_pass = summary_para.add_run(f"✅ Réussis: {passed}  |  ")
    run_pass.font.size = Pt(12)
    run_pass.font.color.rgb = COLOR_PASS
    run_pass.bold = True
    
    run_fail = summary_para.add_run(f"❌ Échoués: {failed}")
    run_fail.font.size = Pt(12)
    run_fail.font.color.rgb = COLOR_FAIL
    run_fail.bold = True

    # Pass rate bar
    rate_para = doc.add_paragraph()
    run_rate = rate_para.add_run(f"Taux de réussite: {pass_rate:.0f}%")
    run_rate.font.size = Pt(14)
    run_rate.bold = True
    run_rate.font.color.rgb = COLOR_PASS if pass_rate >= 70 else COLOR_FAIL

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # DETAILED TEST CASES
    # ══════════════════════════════════════════
    doc.add_heading("Détail des Tests", level=1)

    results = data.get("results", [])
    for i, res in enumerate(results):
        test_id = res.get("test_id", f"TC{i+1:03d}")
        description = res.get("description", "Test sans description")
        status = res.get("status", "unknown")
        error = res.get("error", "")
        screenshot = res.get("screenshot", "")

        # Test case header
        tc_heading = doc.add_heading(f"{test_id}: {description}", level=2)

        # Status line
        status_para = doc.add_paragraph()
        status_para.add_run("Statut: ").bold = True
        if status == "passed":
            run_s = status_para.add_run("✅ RÉUSSI")
            run_s.font.color.rgb = COLOR_PASS
            run_s.bold = True
        elif status == "failed":
            run_s = status_para.add_run("❌ ÉCHOUÉ")
            run_s.font.color.rgb = COLOR_FAIL
            run_s.bold = True
        elif status == "blocked":
            run_s = status_para.add_run("⛔ BLOQUÉ")
            run_s.font.color.rgb = RGBColor(245, 158, 11)
            run_s.bold = True
        else:
            status_para.add_run(status.upper())

        # Error details (if failed)
        if error:
            err_para = doc.add_paragraph()
            err_para.add_run("Raison: ").bold = True
            err_run = err_para.add_run(str(error))
            err_run.font.color.rgb = COLOR_FAIL
            err_run.font.size = Pt(9)

        # Screenshot (if exists)
        if screenshot and os.path.exists(screenshot):
            try:
                doc.add_paragraph()
                doc.add_picture(screenshot, width=Inches(5.0))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = caption.add_run(f"Capture: {os.path.basename(screenshot)}")
                cap_run.font.size = Pt(8)
                cap_run.font.color.rgb = COLOR_MUTED
            except Exception as e:
                doc.add_paragraph(f"[Image non disponible: {screenshot}]")

        doc.add_paragraph()  # Spacing

    # ══════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer.add_run("— Rapport généré par OmniTest Agent —")
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = COLOR_MUTED
    run_footer.italic = True

    # ── Serialize to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def save_report(data: dict) -> str:
    """
    Generate and save a DOCX report to the /reports/ folder.
    Returns the file path.
    """
    os.makedirs("reports", exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"reports/omnitest_report_{ts}.docx"

    doc_bytes = create_test_report_docx(data)
    with open(filename, "wb") as f:
        f.write(doc_bytes)
    
    print(f"[DOCX] Report saved to {filename}")
    return filename


def _browser_label(browser_type: str) -> str:
    """Map internal browser type to human-readable label."""
    labels = {
        "chromium": "Google Chrome",
        "msedge": "Microsoft Edge",
        "firefox": "Mozilla Firefox",
    }
    return labels.get(browser_type, browser_type)
