"""
RF Generator — DOCX Report Generator

Generates a professional Word document report from Robot Framework execution results.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
from pathlib import Path
from io import BytesIO


def generate_rf_docx_report(results: dict, rf_code: str,
                             test_cases: list, base_url: str) -> bytes:
    """
    Generate a DOCX report from RF execution results.

    Args:
        results: Dict with total, passed, failed, failed_tests, test_name, etc.
        rf_code: The generated Robot Framework code string.
        test_cases: List of parsed test case dicts.
        base_url: The base URL of the application under test.

    Returns:
        Bytes of the generated .docx file.
    """
    doc = Document()

    # ── Title ──
    title = doc.add_heading('Rapport de Tests Robot Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Summary Section ──
    doc.add_heading('Résumé Exécution', level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    rows_data = [
        ('Application testée', base_url),
        ('Date exécution', datetime.now().strftime('%d/%m/%Y %H:%M')),
        ('Total tests', str(results.get('total', 0))),
        ('Résultat', f"✅ {results.get('passed', 0)} passés | ❌ {results.get('failed', 0)} échoués"),
    ]
    for i, (key, val) in enumerate(rows_data):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = val

    doc.add_paragraph()

    # ── Test Cases Detail ──
    doc.add_heading('Détail des Cas de Test', level=1)

    failed_tests = results.get('failed_tests', [])

    for tc in test_cases:
        tc_id = tc.get('id', '')
        tc_title = tc.get('title', '')

        # Determine status
        is_failed = any(tc_id in ft or tc_title in ft for ft in failed_tests)
        status = '❌ ÉCHOUÉ' if is_failed else '✅ PASSÉ'

        # TC heading
        doc.add_heading(f'{tc_id} — {tc_title}', level=2)

        # Status
        p = doc.add_paragraph()
        run = p.add_run(f'Statut : {status}')
        run.bold = True

        # Steps
        doc.add_paragraph('Étapes :', style='Intense Quote')
        for i, step in enumerate(tc.get('steps', []), 1):
            if isinstance(step, dict):
                doc.add_paragraph(
                    f"{i}. {step.get('action', '')} → {step.get('expected', '')}",
                    style='List Number'
                )
            else:
                doc.add_paragraph(f"{i}. {step}", style='List Number')

        # Expected result
        doc.add_paragraph(f"Résultat attendu : {tc.get('expected', '')}")

        # Screenshot if exists
        test_name = results.get('test_name', '')
        screenshot_path = f"output/rf_reports/{test_name}/screenshots/{tc_id}.png"
        if os.path.exists(screenshot_path):
            doc.add_picture(screenshot_path, width=Inches(5))

        # Also check for step screenshots
        screenshots_dir = f"output/rf_reports/{test_name}/screenshots"
        if os.path.isdir(screenshots_dir):
            for fname in sorted(os.listdir(screenshots_dir)):
                if fname.startswith(tc_id) and fname.endswith('.png'):
                    full_path = os.path.join(screenshots_dir, fname)
                    if os.path.getsize(full_path) > 0:  # Skip empty screenshots
                        try:
                            doc.add_paragraph(f"📸 {fname}")
                            doc.add_picture(full_path, width=Inches(5))
                        except Exception:
                            pass

        doc.add_paragraph()

    # ── RF Code Section ──
    doc.add_heading('Code Robot Framework Généré', level=1)
    code_para = doc.add_paragraph(rf_code)
    code_para.style = 'No Spacing'
    for run in code_para.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    # ── Save to bytes ──
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
